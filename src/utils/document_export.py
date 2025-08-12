from typing import Dict, List, Optional, Union
import json
from pathlib import Path
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime
from PyQt6.QtCore import QThread, pyqtSignal
import logging
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import arabic_reshaper
from bidi.algorithm import get_display

class DocumentExporter(QThread):
    """Export OCR results in various formats"""
    
    progress = pyqtSignal(int)
    status = pyqtSignal(str)
    
    def __init__(self, output_dir: str):
        super().__init__()
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._setup_fonts()
        
    def _setup_fonts(self):
        """Set up fonts for PDF generation with Arabic support"""
        try:
            # Register Arabic font
            pdfmetrics.registerFont(TTFont('Arabic', 'assets/fonts/NotoSansArabic-Regular.ttf'))
        except Exception as e:
            logging.error(f"Failed to register Arabic font: {str(e)}")
    
    def export_to_pdf(self, 
                     results: Dict,
                     template_info: Dict,
                     output_file: str) -> str:
        """Export OCR results to PDF with proper formatting"""
        try:
            pdf_path = self.output_dir / f"{output_file}.pdf"
            c = canvas.Canvas(str(pdf_path))
            
            # Set up document
            c.setTitle(template_info['name'])
            
            # Add header
            self._add_pdf_header(c, template_info)
            
            # Add content
            y_position = 750
            for lang in ['french', 'arabic']:
                if lang in results:
                    for result in results[lang]:
                        text = result.text
                        if lang == 'arabic':
                            text = get_display(arabic_reshaper.reshape(text))
                            c.setFont('Arabic', 12)
                        else:
                            c.setFont('Helvetica', 12)
                        
                        c.drawString(72, y_position, text)
                        y_position -= 20
            
            c.save()
            return str(pdf_path)
            
        except Exception as e:
            logging.error(f"PDF export failed: {str(e)}")
            raise
    
    def export_to_word(self,
                      results: Dict,
                      template_info: Dict,
                      output_file: str) -> str:
        """Export OCR results to Word document"""
        try:
            doc = docx.Document()
            
            # Add header
            header = doc.add_heading(template_info['name'], 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add Arabic header
            ar_header = doc.add_paragraph()
            ar_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ar_run = ar_header.add_run(template_info['name_ar'])
            ar_run.font.size = Pt(16)
            
            # Add content
            for lang in ['french', 'arabic']:
                if lang in results:
                    doc.add_heading(f"{lang.title()} Text", 1)
                    for result in results[lang]:
                        p = doc.add_paragraph()
                        p.add_run(result.text)
                        
            # Save document
            word_path = self.output_dir / f"{output_file}.docx"
            doc.save(str(word_path))
            return str(word_path)
            
        except Exception as e:
            logging.error(f"Word export failed: {str(e)}")
            raise
    
    def export_to_excel(self,
                       results: Dict,
                       template_info: Dict,
                       output_file: str) -> str:
        """Export OCR results to Excel spreadsheet"""
        try:
            # Prepare data for Excel
            data = []
            for lang in ['french', 'arabic']:
                if lang in results:
                    for result in results[lang]:
                        data.append({
                            'Language': lang,
                            'Text': result.text,
                            'Confidence': result.confidence,
                            'X': result.bounding_box[0],
                            'Y': result.bounding_box[1],
                            'Width': result.bounding_box[2],
                            'Height': result.bounding_box[3]
                        })
            
            # Create DataFrame and export
            df = pd.DataFrame(data)
            excel_path = self.output_dir / f"{output_file}.xlsx"
            df.to_excel(str(excel_path), index=False)
            return str(excel_path)
            
        except Exception as e:
            logging.error(f"Excel export failed: {str(e)}")
            raise
    
    def export_to_json(self,
                      results: Dict,
                      template_info: Dict,
                      output_file: str) -> str:
        """Export OCR results to structured JSON"""
        try:
            # Prepare structured data
            export_data = {
                'metadata': {
                    'template_name': template_info['name'],
                    'template_name_ar': template_info['name_ar'],
                    'export_date': datetime.now().isoformat(),
                    'template_version': template_info.get('template_version', '1.0')
                },
                'results': {
                    'french': [
                        {
                            'text': r.text,
                            'confidence': r.confidence,
                            'position': r.bounding_box
                        }
                        for r in results.get('french', [])
                    ],
                    'arabic': [
                        {
                            'text': r.text,
                            'confidence': r.confidence,
                            'position': r.bounding_box
                        }
                        for r in results.get('arabic', [])
                    ]
                }
            }
            
            # Save JSON file
            json_path = self.output_dir / f"{output_file}.json"
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, ensure_ascii=False, indent=2)
                
            return str(json_path)
            
        except Exception as e:
            logging.error(f"JSON export failed: {str(e)}")
            raise
    
    def batch_export(self,
                    results: Dict[str, Dict],
                    template_info: Dict,
                    formats: List[str]) -> Dict[str, List[str]]:
        """Export multiple documents in various formats"""
        exported_files = {fmt: [] for fmt in formats}
        total_exports = len(results) * len(formats)
        completed = 0
        
        try:
            for doc_path, doc_results in results.items():
                base_name = Path(doc_path).stem
                
                for fmt in formats:
                    self.status.emit(f"Exporting {base_name} to {fmt}")
                    
                    export_method = getattr(self, f"export_to_{fmt}")
                    output_path = export_method(
                        doc_results.get('ocr_results', {}),
                        template_info,
                        base_name
                    )
                    exported_files[fmt].append(output_path)
                    
                    completed += 1
                    self.progress.emit(int(completed / total_exports * 100))
                    
            return exported_files
            
        except Exception as e:
            logging.error(f"Batch export failed: {str(e)}")
            raise
